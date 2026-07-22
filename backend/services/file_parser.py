"""多格式文件解析器：docx / pdf / txt / md / xlsx"""
import io
import os
from typing import Optional


class FileParser:
    """统一的文件解析接口"""

    @staticmethod
    async def parse(file_path: str, file_type: str) -> str:
        """根据文件类型解析文件内容"""
        file_type = file_type.lower().lstrip(".")

        parsers = {
            "docx": FileParser._parse_docx,
            "pdf": FileParser._parse_pdf,
            "txt": FileParser._parse_txt,
            "md": FileParser._parse_txt,
            "xlsx": FileParser._parse_xlsx,
        }

        parser = parsers.get(file_type)
        if parser is None:
            raise ValueError(f"不支持的文件类型: {file_type}")

        return parser(file_path)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析 Word 文档"""
        from docx import Document

        doc = Document(file_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 识别标题层级
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                try:
                    level_num = int(level)
                    prefix = "#" * level_num
                    lines.append(f"{prefix} {text}")
                except ValueError:
                    lines.append(text)
            else:
                lines.append(text)

        # 解析表格
        for table in doc.tables:
            lines.append("")  # 空行分隔
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析 PDF 文件"""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        lines = []

        for page in doc:
            text = page.get_text("text")
            if text:
                lines.append(text.strip())

        doc.close()
        return "\n\n".join(lines)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """解析纯文本 / Markdown 文件"""
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # 最后尝试忽略错误
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _parse_xlsx(file_path: str) -> str:
        """解析 Excel 文件"""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"## 工作表: {sheet_name}")
            lines.append("")

            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cells):  # 跳过全空行
                    lines.append("| " + " | ".join(cells) + " |")

            lines.append("")

        wb.close()
        return "\n".join(lines)

    @staticmethod
    def extract_docx_toc(file_path: str) -> list[dict]:
        """从 Word 文档提取目录结构（基于 Heading 样式）"""
        from docx import Document

        doc = Document(file_path)
        chapters = []
        stack = []  # 用于追踪层级关系

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                except ValueError:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                node = {
                    "title": text,
                    "level": level,
                    "title_only": False,
                    "content_type": "text",
                    "content_prompt": "",
                    "children": [],
                }

                # 找到父节点
                while stack and stack[-1]["level"] >= level:
                    stack.pop()

                if stack:
                    stack[-1]["children"].append(node)
                else:
                    chapters.append(node)

                stack.append(node)

        return chapters

    @staticmethod
    def parse_from_bytes(data: bytes, file_type: str) -> str:
        """从字节数据解析文件内容"""
        file_type = file_type.lower().lstrip(".")

        if file_type == "docx":
            from docx import Document
            doc = Document(io.BytesIO(data))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            return "\n".join(lines)

        elif file_type == "pdf":
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            lines = []
            for page in doc:
                text = page.get_text("text")
                if text:
                    lines.append(text.strip())
            doc.close()
            return "\n\n".join(lines)

        elif file_type in ("txt", "md"):
            for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    return data.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return data.decode("utf-8", errors="ignore")

        elif file_type == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"## 工作表: {sheet_name}")
                lines.append("")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    if any(cells):
                        lines.append("| " + " | ".join(cells) + " |")
                lines.append("")
            wb.close()
            return "\n".join(lines)

        else:
            raise ValueError(f"不支持的文件类型: {file_type}")