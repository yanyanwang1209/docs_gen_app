"""MD 转 Word 转换服务"""
import io
import re
import os
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class Md2WordConverter:
    """将 Markdown 文本转换为 Word 文档"""

    # 字体配置
    FONT_NAME = "宋体"
    CODE_FONT = "Consolas"
    HEADING_COLOR = RGBColor(0, 0, 0)  # 标题黑色

    # 字号配置（单位：Pt）
    DOC_TITLE_SIZE = Pt(22)       # 文档标题：二号
    H1_SIZE = Pt(16)             # 一级标题：三号
    H2_SIZE = Pt(15)             # 二级标题：小三
    H3_SIZE = Pt(14)             # 三级标题：四号
    H4_SIZE = Pt(12)             # 四级标题：小四
    BODY_SIZE = Pt(12)           # 正文：小四
    TABLE_TEXT_SIZE = Pt(10.5)   # 表格文字：五号
    LINE_SPACING = 1.5           # 行间距

    def __init__(self, doc_title: Optional[str] = None):
        self.doc = Document()
        self.doc_title = doc_title  # 外部传入的文档标题
        self._setup_styles()
        self._first_h1_used = False

    def _setup_styles(self):
        """预配置文档样式"""
        # 设置默认字体
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.FONT_NAME
        font.size = self.BODY_SIZE
        font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)

        # 段落格式
        pf = style.paragraph_format
        pf.line_spacing = self.LINE_SPACING
        pf.first_line_indent = Cm(0.74)  # 约 2 个中文字符
        pf.space_after = Pt(6)

    def _add_heading(self, text: str, level: int):
        """添加标题（黑色字体）"""
        # 使用内置 Heading 样式
        heading = self.doc.add_heading(text, level=level)

        # 设置字体和颜色
        for run in heading.runs:
            run.font.name = self.FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
            run.font.bold = True
            run.font.color.rgb = self.HEADING_COLOR

            if level == 0:  # 文档标题
                run.font.size = self.DOC_TITLE_SIZE
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 1:
                run.font.size = self.H1_SIZE
            elif level == 2:
                run.font.size = self.H2_SIZE
            elif level == 3:
                run.font.size = self.H3_SIZE
            else:
                run.font.size = self.H4_SIZE

        # 标题不缩进
        heading.paragraph_format.first_line_indent = Pt(0)

    def _add_paragraph(self, text: str):
        """添加正文段落，支持 **加粗** 行内标记"""
        if not text.strip():
            return

        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing = self.LINE_SPACING
        pf.first_line_indent = Cm(0.74)

        # 解析 **加粗** 行内标记
        self._add_inline_runs(para, text, self.BODY_SIZE)

    def _add_inline_runs(self, para, text: str, base_size: Pt):
        """解析文本中的 **加粗** 和 <br> 换行标记，添加对应 run"""
        # 先按 <br> 分割处理换行
        segments = re.split(r"<br\s*/?>", text)
        for seg_idx, segment in enumerate(segments):
            if seg_idx > 0:
                # 在段落内添加换行
                run = para.add_run()
                run.add_break()

            # 匹配 **text** 或 __text__
            pattern = re.compile(r"(\*\*(.+?)\*\*|__(.+?)__)")
            last_end = 0
            for match in pattern.finditer(segment):
                # 前面的普通文本
                before = segment[last_end:match.start()]
                if before:
                    run = para.add_run(before)
                    run.font.name = self.FONT_NAME
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
                    run.font.size = base_size
                    run.font.color.rgb = RGBColor(0, 0, 0)

                # 加粗文本
                bold_text = match.group(2) or match.group(3)
                run = para.add_run(bold_text)
                run.font.name = self.FONT_NAME
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
                run.font.size = base_size
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

                last_end = match.end()

            # 剩余文本
            remaining = segment[last_end:]
            if remaining:
                run = para.add_run(remaining)
                run.font.name = self.FONT_NAME
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
                run.font.size = base_size
                run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_code_block(self, code_lines: list):
        """添加代码块"""
        code_text = "\n".join(code_lines)

        # 添加浅灰背景的段落
        para = self.doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = self.CODE_FONT
        run.font.size = self.TABLE_TEXT_SIZE

        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.line_spacing = 1.0

    def _add_table(self, table_lines: list):
        """添加表格，自动合并行尾空列"""
        if not table_lines:
            return

        # 解析表格行，过滤掉表头分隔行（| --- | :--- | :---: | 等）
        rows = []
        for line in table_lines:
            line = line.strip().strip("|")
            cells = [cell.strip() for cell in line.split("|")]
            # 跳过表头分隔行
            if all(re.match(r"^:?-{3,}:?$", c) for c in cells):
                continue
            rows.append(cells)

        if not rows:
            return

        # 创建 Word 表格
        max_cols = max(len(row) for row in rows)
        if max_cols == 0:
            return

        # 将所有行补齐到 max_cols（缺失的列视为空）
        for row_data in rows:
            while len(row_data) < max_cols:
                row_data.append("")

        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    cell = table.cell(i, j)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.paragraph_format.first_line_indent = Pt(0)
                    # 解析 **加粗** 标记
                    self._add_inline_runs(para, cell_text, self.TABLE_TEXT_SIZE)

        # 自动合并：从左往右扫描，每个有值列与其后连续空列合并
        for i, row_data in enumerate(rows):
            j = 0
            while j < max_cols:
                if row_data[j].strip():
                    # 该列有值，向后查找连续空列
                    k = j + 1
                    while k < max_cols and not row_data[k].strip():
                        k += 1
                    # 如果有连续空列，合并后清理多余空段落
                    if k > j + 1:
                        table.cell(i, j).merge(table.cell(i, k - 1))
                        # 清理合并后单元格中多余的空段落，只保留有内容的段落
                        merged_cell = table.cell(i, j)
                        merged_paras = merged_cell.paragraphs
                        # 从后往前删除空白段落，至少保留一个
                        for p_idx in range(len(merged_paras) - 1, 0, -1):
                            p = merged_paras[p_idx]
                            if not p.text.strip():
                                p._element.getparent().remove(p._element)
                    j = k
                else:
                    j += 1

        # 加粗表头：2列表加粗第一列（字段名），多列表加粗第一行
        if max_cols == 2:
            for row in table.rows:
                for para in row.cells[0].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
        else:
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True

        # 表格后添加一个空段落分隔，避免两个表格在 Word 中粘在一起
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.line_spacing = self.LINE_SPACING
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.first_line_indent = Pt(0)

    def _add_image_placeholder(self, alt_text: str):
        """添加图片占位符"""
        para = self.doc.add_paragraph()
        run = para.add_run(f"[图片: {alt_text}]")
        run.font.name = self.FONT_NAME
        run.font.size = Pt(9)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_list_item(self, text: str, indent: int = 0):
        """添加列表项，支持 **加粗**"""
        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing = self.LINE_SPACING
        pf.left_indent = Cm(0.74 * (indent + 1))
        pf.first_line_indent = Pt(0)

        self._add_inline_runs(para, text, self.BODY_SIZE)

    def convert(self, markdown_text: str) -> Document:
        """将 Markdown 文本转换为 Word 文档"""
        # 如果传入了文档标题，先添加文档标题
        if self.doc_title:
            self._add_heading(self.doc_title, level=0)
            self._first_h1_used = True  # 标记已添加，避免第一个 H1 被当作标题

        lines = markdown_text.split("\n")
        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []

        while i < len(lines):
            line = lines[i]

            # 代码块处理
            if line.strip().startswith("```"):
                if in_code_block:
                    self._add_code_block(code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 表格处理
            if line.strip().startswith("|") and line.strip().endswith("|"):
                if not in_table:
                    in_table = True
                table_lines.append(line)
                i += 1
                continue
            elif in_table:
                if line.strip().startswith("|"):
                    table_lines.append(line)
                    i += 1
                    continue
                else:
                    self._add_table(table_lines)
                    table_lines = []
                    in_table = False

            # 标题处理
            heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(text, level=level)
                i += 1
                continue

            # 列表项
            list_match = re.match(r"^(\s*)[-*+]\s+(.+)", line)
            if list_match:
                indent = len(list_match.group(1)) // 2
                text = list_match.group(2)
                self._add_list_item(text, indent)
                i += 1
                continue

            # 有序列表
            ordered_match = re.match(r"^(\s*)\d+\.\s+(.+)", line)
            if ordered_match:
                indent = len(ordered_match.group(1)) // 2
                text = ordered_match.group(2)
                self._add_list_item(text, indent)
                i += 1
                continue

            # 图片
            image_match = re.match(r"!\[(.*?)\]\((.+?)\)", line)
            if image_match:
                alt_text = image_match.group(1) or "图片"
                self._add_image_placeholder(alt_text)
                i += 1
                continue

            # 空行
            if not line.strip():
                i += 1
                continue

            # 普通段落
            self._add_paragraph(line)
            i += 1

        # 处理剩余的表格
        if in_table and table_lines:
            self._add_table(table_lines)

        return self.doc

    def save(self, output_path: str):
        """保存 Word 文档"""
        self.doc.save(output_path)

    def save_to_bytes(self) -> bytes:
        """保存为字节流"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()